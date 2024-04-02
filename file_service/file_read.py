from os.path import join, dirname, exists
from qrcode import QRCode, constants
from docx.shared import Inches, Pt
from os.path import join, dirname
from datetime import datetime
from docx2pdf import convert
from docx import Document
from os import remove


async def replace_text(paragraph, old_text, new_text, font_size=None, bold=True, underline=False):
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
            if underline:
                run.font.underline = True


async def replace_table_text(table, old_text, new_text, font_size=None, bold=True, size=1.6):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if old_text in run.text:
                        if old_text == "&":
                            run.add_picture(join(dirname(__file__), f"file_qrcode/{new_text}.png"),
                                            width=Inches(size))
                        if old_text != "&":
                            run.text = run.text.replace(old_text, new_text)
                        if font_size:
                            run.font.size = Pt(font_size)
                        if bold:
                            run.font.bold = True


async def process_document(address, name):
    doc = Document(join(dirname(__file__), "ariza.docx"))
    for paragraph in doc.paragraphs:
        if "ADDRESS" in paragraph.text:
            await replace_text(paragraph, "ADDRESS", address)
        if "NAME" in paragraph.text:
            await replace_text(paragraph, "NAME", name)
        if "DATEFULL" in paragraph.text:
            await replace_text(paragraph, "DATEFULL", f"{datetime.now().strftime('%d.%m.%Y')}    {name}")
        doc.save(join(dirname(__file__), f"file_ariza\\{name}.docx"))


async def process_contract(name, faculty, passport, number, address):
    doc = Document(join(dirname(__file__), "cotract.docx"))
    for paragraph in doc.paragraphs:
        if "SONLI1" in paragraph.text:
            await replace_text(paragraph, "SONLI1", f"0001")
        if "DATE" in paragraph.text:
            await replace_text(paragraph, "DATE", f"{datetime.now().strftime('%d.%m.')}", bold=False)
        if "NAME" in paragraph.text:
            await replace_text(paragraph, "NAME", f"{name}")
        if "FACULTY" in paragraph.text:
            await replace_text(paragraph, "FACULTY", f"{faculty}")
        if "SS!" in paragraph.text:
            await replace_text(paragraph, "SER@", f"{passport}")
        if "ADDRES" in paragraph.text:
            await replace_text(paragraph, "PASSR", f"{address}")
        if "TELNUMBER" in paragraph.text:
            await replace_text(paragraph, "TELNUMBER", f"{number}")
    for table in doc.tables:
        await replace_table_text(table=table, old_text="RFS", new_text=passport, font_size=12)
        await replace_table_text(table=table, old_text="FISH", new_text=name, font_size=12)
        await replace_table_text(table=table, old_text="TELNUMBER", new_text=number, font_size=12)
        await replace_table_text(table=table, old_text="ADDRES", new_text=address, font_size=12)
        await replace_table_text(table=table, old_text="&", new_text=name, bold=False)
    doc.save(join(dirname(__file__), f"file_shartnoma\\{name}.docx"))


async def func_qrcode(url, name):
    qr = QRCode(
        version=1,
        error_correction=constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )

    qr.add_data(f"https://swine-viable-luckily.ngrok-free.app/get_file/{url}")
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    return img.save(join(dirname(__file__), f"file_qrcode/{name}.png"))
